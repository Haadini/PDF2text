from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import tempfile
from docx import Document

pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

def remove_non_xml_chars(text, type):
    """
    Remove non-XML-compatible characters from the text.
    """
    return re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\xff]', '', text)
def pdf_to_word(pdf_path, output_dir, lang='fas', **kwargs):
    """ 
    A way to use google-tesseract ocr for extracting 
    texts from pdf file.
    
    Args:
        pdf_path (str): PDF file path.
        output_dir (str): Output directory.
        lang (str): tesseract language support.
    """
    
    pdf_name = pdf_path.split('/')[-1].split('.')[0]
    pages = convert_from_path(pdf_path, **kwargs)
    texts = []
    
    print(f'PDF is preparing to convert into document [#{len(pages)} pages]')
    for i, page in tqdm(enumerate(pages), position=0):
        
        with tempfile.TemporaryDirectory() as img_dir:
            img_name = f'{pdf_name}-{i+1}.jpg'
            img_path = os.path.join(img_dir, img_name)
            
            page.save(img_path, 'JPEG')
            text = pytesseract.image_to_string(Image.open(img_path), lang=lang)
            cleaned_text = remove_non_xml_chars(text)
            texts.append(cleaned_text)
    
    document = Document()
    style_normal = document.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.rtl = True
    
    style_h1 = document.styles['Heading 1']
    font = style_h1.font
    font.name = 'Arial'
    font.rtl = True
    
    for i, text in tqdm(enumerate(texts), position=0):
        heading = document.add_heading(f'صفحه: {i+1}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.style = document.styles['Heading 1']
        
        paragraph = document.add_paragraph(text.encode('utf-8').decode('utf-8'))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.style = document.styles['Normal']
    
    output_path = os.path.join(output_dir, f'{pdf_name}.{type}')
    document.save(output_path)
    print(f'Done! Your document can be find here "{output_path}"')

    doc = Document(output_path)
    text = ""
    
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    output_file = os.path.join(output_dir, f'{pdf_name}.txt')
    with open(output_file, "w", encoding="utf-8") as file:
      file.write(text)
    